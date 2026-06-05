from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from config import settings
from rag.db import db
from services.orchestrator import orchestrator
import fitz
import json
import random
from collections import Counter
from rag.embedder import embedder
from rag.in_memory_db import memory_db

app = FastAPI(title="Study Planner Backend", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Lifecycle ──────────────────────────────────────────────────────────────────

@app.on_event("startup")
async def startup_event():
    await db.connect()

@app.on_event("shutdown")
async def shutdown_event():
    await db.disconnect()


# ── Health ─────────────────────────────────────────────────────────────────────

@app.get("/health")
def health_check():
    return {"status": "healthy"}


# ── Documents ──────────────────────────────────────────────────────────────────

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    content = await file.read()
    text = ""
    filename_lower = file.filename.lower()
    if filename_lower.endswith(".pdf"):
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
    elif filename_lower.endswith(".docx"):
        import io
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        doc = DocxDocument(io.BytesIO(content))
        parts: list[str] = []

        def extract_text_from_element(el) -> str:
            """Recursively collect all <w:t> text runs from an XML element."""
            return "".join(node.text or "" for node in el.iter(qn("w:t")))

        # Paragraphs (main body, headers, footers)
        for para in doc.paragraphs:
            line = para.text.strip()
            if line:
                parts.append(line)

        # Tables — iterate every cell
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)

        # Headers and footers
        for section in doc.sections:
            for hdr_ftr in [section.header, section.footer,
                            section.even_page_header, section.even_page_footer,
                            section.first_page_header, section.first_page_footer]:
                try:
                    for para in hdr_ftr.paragraphs:
                        line = para.text.strip()
                        if line:
                            parts.append(line)
                except Exception:
                    pass

        # Text boxes and other drawing elements via raw XML
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            raw = extract_text_from_element(txbx).strip()
            if raw:
                parts.append(raw)

        text = "\n".join(parts)
    else:
        text = content.decode("utf-8", errors="ignore")

    text = text.replace("\x00", "")

    chunk_size = 300
    words = text.split()
    chunks = [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]

    chunk_index = 0
    for c in chunks:
        c = c.replace("\x00", "")
        if not c.strip():
            continue
        emb = embedder.embed(c)
        if db.pool:
            async with db.pool.acquire() as conn:
                await conn.execute(
                    """
                    INSERT INTO legal_chunks (source_file, chunk_index, content, embedding, doc_type)
                    VALUES ($1, $2, $3, $4::vector, $5)
                    """,
                    file.filename, chunk_index, c, json.dumps(emb), "study_material",
                )
        else:
            memory_db.add_chunk(file.filename, c, emb)
        chunk_index += 1

    return {"status": "success", "filename": file.filename, "chunks_added": chunk_index}


@app.get("/documents")
async def list_documents():
    """Return all indexed documents with chunk counts."""
    if db.pool:
        async with db.pool.acquire() as conn:
            rows = await conn.fetch(
                """
                SELECT source_file,
                       COUNT(*) AS chunk_count,
                       MAX(created_at) AS indexed_at
                FROM legal_chunks
                GROUP BY source_file
                ORDER BY MAX(created_at) DESC
                """
            )
            return {
                "documents": [
                    {
                        "filename": r["source_file"],
                        "chunks": r["chunk_count"],
                        "indexed_at": r["indexed_at"].isoformat() if r["indexed_at"] else None,
                    }
                    for r in rows
                ]
            }
    else:
        counts = Counter(c["source_file"] for c in memory_db.chunks)
        return {
            "documents": [
                {"filename": name, "chunks": count, "indexed_at": None}
                for name, count in counts.items()
            ]
        }


@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    """Remove a document and all its chunks."""
    if db.pool:
        async with db.pool.acquire() as conn:
            await conn.execute(
                "DELETE FROM legal_chunks WHERE source_file = $1", filename
            )
        return {"status": "success", "deleted": filename}
    else:
        indices_to_keep = [
            i for i, c in enumerate(memory_db.chunks)
            if c["source_file"] != filename
        ]
        memory_db.chunks = [memory_db.chunks[i] for i in indices_to_keep]
        memory_db.embeddings = [memory_db.embeddings[i] for i in indices_to_keep]
        return {"status": "success", "deleted": filename}


# ── Chat ───────────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    top_k: int = 5
    similarity_threshold: float = 0.0
    temperature: float = 0.3
    mode: str = "ai"  # "ai" | "corpus" | "compare"
    source_file: str | None = None  # if set, restrict to this document


@app.post("/chat")
async def chat_with_notes(req: ChatRequest):
    kwargs = dict(
        top_k=req.top_k,
        similarity_threshold=req.similarity_threshold,
        temperature=req.temperature,
        source_file=req.source_file,
    )
    if req.mode == "corpus":
        result = await orchestrator.corpus_only(req.message, **kwargs)
    elif req.mode == "compare":
        result = await orchestrator.compare(req.message, **kwargs)
    else:
        result = await orchestrator.chat(req.message, **kwargs)
    return result


# ── Quiz ───────────────────────────────────────────────────────────────────────

class QuizGenerateRequest(BaseModel):
    num_questions: int = 5
    source_file: str | None = None  # if set, restrict to this document


@app.post("/quiz/generate")
async def generate_quiz(req: QuizGenerateRequest):
    if db.pool:
        async with db.pool.acquire() as conn:
            if req.source_file:
                rows = await conn.fetch(
                    "SELECT content FROM legal_chunks WHERE source_file = $1 ORDER BY random() LIMIT 20",
                    req.source_file,
                )
            else:
                rows = await conn.fetch(
                    "SELECT content FROM legal_chunks ORDER BY random() LIMIT 20"
                )
            sample_chunks = [{"content": r["content"]} for r in rows]
    else:
        pool = [c for c in memory_db.chunks if not req.source_file or c["source_file"] == req.source_file]
        sample_chunks = random.sample(pool, min(20, len(pool)))

    if not sample_chunks:
        return {"status": "error", "message": "No documents uploaded yet. Please upload a document first."}

    # Filter out noisy / very short chunks before sending to the model
    def is_meaningful(text: str) -> bool:
        stripped = text.strip()
        # Skip chunks that are too short, look like filenames/metadata, or are mostly non-alpha
        if len(stripped) < 80:
            return False
        alpha_ratio = sum(c.isalpha() or c.isspace() for c in stripped) / max(len(stripped), 1)
        if alpha_ratio < 0.5:
            return False
        return True

    meaningful_chunks = [c for c in sample_chunks if is_meaningful(c["content"])]
    # Fall back to all chunks if filtering removed everything
    if not meaningful_chunks:
        meaningful_chunks = sample_chunks

    context = "\n\n---\n\n".join([c["content"] for c in meaningful_chunks])

    prompt = (
        f"You are a university-level study assistant helping a student prepare for an exam.\n"
        f"Below are passages extracted from a student's uploaded document.\n\n"
        f"Your task: Generate exactly {req.num_questions} meaningful study questions based SOLELY on the "
        f"CONCEPTS, FACTS, DEFINITIONS, PROCESSES, or ARGUMENTS present in the passages below.\n\n"
        f"STRICT RULES:\n"
        f"- Questions MUST test understanding of the subject matter (concepts, definitions, processes, arguments, comparisons)\n"
        f"- Do NOT ask about the document itself, its filename, file format, word count, structure, or metadata\n"
        f"- Do NOT ask 'what is discussed in this document' or similar meta-questions\n"
        f"- Every answer MUST be directly supported by information in the passages\n"
        f"- Use a mix of difficulty levels: Easy (recall), Medium (comprehension), Hard (analysis/application)\n\n"
        f"Return ONLY a valid JSON array. No explanation, no markdown, no extra text outside the array.\n"
        f"Each element: {{\"q\": \"<question>\", \"a\": \"<detailed answer from the text>\", \"diff\": \"Easy|Medium|Hard\"}}\n\n"
        f"Passages:\n{context}"
    )

    if not orchestrator.client:
        return {"status": "error", "message": "MISTRAL_API_KEY is missing in backend/.env"}

    try:
        response = await orchestrator.client.chat.complete_async(
            model=settings.MISTRAL_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        text = response.choices[0].message.content.strip()

        # Strip any markdown fences the model may have added
        if text.startswith("```"):
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
            text = text.strip()

        # Extract the JSON array even if the model appended extra prose
        start = text.find("[")
        end = text.rfind("]")
        if start == -1 or end == -1 or end < start:
            return {"status": "error", "message": "Model did not return a valid JSON array. Please try again."}
        text = text[start:end + 1]

        quiz_data = json.loads(text)

        # Normalise — ensure every item has the expected keys
        cleaned = []
        for item in quiz_data:
            if isinstance(item, dict) and "q" in item and "a" in item:
                cleaned.append({
                    "q": str(item.get("q", "")),
                    "a": str(item.get("a", "")),
                    "diff": str(item.get("diff", "Medium")),
                })
        if not cleaned:
            return {"status": "error", "message": "Model returned an empty quiz. Please try again."}

        return {"status": "success", "quiz": cleaned}
    except json.JSONDecodeError as e:
        return {"status": "error", "message": f"Could not parse quiz JSON: {e}. Please try again."}
    except Exception as e:
        return {"status": "error", "message": f"Error generating quiz: {e}"}
