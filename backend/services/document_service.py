import asyncio
import io
import json

import fitz
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from fastapi import UploadFile

from rag.db import db


class DocumentService:
    async def upload_document(self, file: UploadFile):
        from rag.embedder import embedder

        content = await file.read()
        chunks = await asyncio.to_thread(
            self._extract_and_chunk,
            file.filename,
            content,
        )
        chunks = [
            chunk.replace("\x00", "")
            for chunk in chunks
            if chunk.replace("\x00", "").strip()
        ]
        embeddings = await embedder.embed_many(chunks)
        await self._store_chunks(file.filename, chunks, embeddings)

        return {
            "status": "success",
            "filename": file.filename,
            "chunks_added": len(chunks),
        }

    def _extract_and_chunk(self, filename: str, content: bytes) -> list[str]:
        text = self._extract_text(filename, content).replace("\x00", "")
        return self._chunk_text(text)

    async def list_documents(self):
        if not db.pool:
            raise RuntimeError("PostgreSQL is not connected.")

        async with db.pool.acquire() as conn:
            rows = await conn.fetch(
                """
                SELECT source_file,
                       COUNT(*) AS chunk_count,
                       MAX(created_at) AS indexed_at
                FROM corpus_chunks
                GROUP BY source_file
                ORDER BY MAX(created_at) DESC
                """
            )
        return {
            "documents": [
                {
                    "filename": row["source_file"],
                    "chunks": row["chunk_count"],
                    "indexed_at": row["indexed_at"].isoformat()
                    if row["indexed_at"]
                    else None,
                }
                for row in rows
            ]
        }

    async def delete_document(self, filename: str):
        if not db.pool:
            raise RuntimeError("PostgreSQL is not connected.")

        async with db.pool.acquire() as conn:
            await conn.execute(
                "DELETE FROM corpus_chunks WHERE source_file = $1",
                filename,
            )
        return {"status": "success", "deleted": filename}

    def _extract_text(self, filename: str, content: bytes) -> str:
        filename_lower = filename.lower()
        if filename_lower.endswith(".pdf"):
            return self._extract_pdf_text(content)
        if filename_lower.endswith(".docx"):
            return self._extract_docx_text(content)
        return content.decode("utf-8", errors="ignore")

    def _extract_pdf_text(self, content: bytes) -> str:
        document = fitz.open(stream=content, filetype="pdf")
        return "\n".join(page.get_text() for page in document)

    def _extract_docx_text(self, content: bytes) -> str:
        document = DocxDocument(io.BytesIO(content))
        parts: list[str] = []

        def extract_text_from_element(element) -> str:
            return "".join(node.text or "" for node in element.iter(qn("w:t")))

        for paragraph in document.paragraphs:
            line = paragraph.text.strip()
            if line:
                parts.append(line)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)

        for section in document.sections:
            headers_and_footers = [
                section.header,
                section.footer,
                section.even_page_header,
                section.even_page_footer,
                section.first_page_header,
                section.first_page_footer,
            ]
            for header_or_footer in headers_and_footers:
                try:
                    for paragraph in header_or_footer.paragraphs:
                        line = paragraph.text.strip()
                        if line:
                            parts.append(line)
                except Exception:
                    pass

        for text_box in document.element.body.iter(qn("w:txbxContent")):
            raw = extract_text_from_element(text_box).strip()
            if raw:
                parts.append(raw)

        return "\n".join(parts)

    def _chunk_text(self, text: str) -> list[str]:
        chunk_size = 300
        words = text.split()
        return [
            " ".join(words[index : index + chunk_size])
            for index in range(0, len(words), chunk_size)
        ]

    async def _store_chunks(
        self,
        filename: str,
        chunks: list[str],
        embeddings: list[list[float]],
    ) -> None:
        if not db.pool or not db.vector_enabled:
            raise RuntimeError("PostgreSQL with pgvector is not available.")

        async with db.pool.acquire() as conn:
            async with conn.transaction():
                await conn.executemany(
                    """
                    INSERT INTO corpus_chunks (
                        source_file, chunk_index, content, embedding, doc_type
                    )
                    VALUES ($1, $2, $3, $4::vector, 'study_material')
                    """,
                    [
                        (filename, index, content, json.dumps(embedding))
                        for index, (content, embedding) in enumerate(
                            zip(chunks, embeddings, strict=True)
                        )
                    ],
                )


document_service = DocumentService()
